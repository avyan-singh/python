import docx
import PyPDF2 as pypdf2
import os
from pathlib import Path
import time
import threading
def search_text_in_directory_directory(text="hydrogen",basedir=Path(r"C:\avyan\python\IMPORTANT_MODULES")):
    text=text.lower()
    basedir=Path(r"C:\avyan\python\IMPORTANT_MODULES")
    dirs=[os.path.join(basedir,d) for d in os.listdir(basedir) if os.path.isdir(os.path.join(basedir,d))]
    for d in basedir.iterdir():
        if d.is_dir():
            subdir=d
            for di in subdir.iterdir():
                if di.is_dir():
                    for name in os.listdir(di):
                        path=os.path.join(di,name).lower()
                        if path.endswith('.txt'):
                            with open(path,'r') as f:
                                data=f.read()
                                data=data.lower()
                                if text in data:
                                    print(f"Found in text file: {path}")
                                    break                       
            for name in os.listdir(subdir):
                path=os.path.join(subdir,name).lower()
                if path.endswith('.txt'):
                    with open(path,'r') as f:
                        data=f.read()
                        data=data.lower()
                        if text in data:
                            print(f"Found in text file: {path}")
                            break
                elif path.endswith('.pdf'):
                    with open(path,'rb') as pdf:
                        pdf_reader=pypdf2.PdfReader(pdf)
                        num=-1
                        for page in pdf_reader.pages:
                            num=num+1
                            pdftext=pdf_reader.pages[num].extract_text()
                            pdftext=pdftext.lower()
                            if text in str(pdftext):
                                print(f"Found in PDF file: {path}")
                                break
                elif path.endswith('.docx'):
                    with open(path,'rb') as docx_file:
                        document_reader=docx.Document(docx_file)
                        fulltext=[]
                        for para in document_reader.paragraphs:
                            fulltext.append(para.text)
                            docx_text='\n'.join(fulltext)
                            docx_text=docx_text.lower()
                        if text in docx_text:
                            print(f"Found in DOCX file: {path}")
                            break
def search_text_in_directory(text="hydrogen",basedir=Path(r"C:\avyan\python\IMPORTANT_MODULES")):
    text=text.lower()
    basedir=Path(r"C:\avyan\python\IMPORTANT_MODULES")
    dirs=[os.path.join(basedir,d) for d in os.listdir(basedir) if os.path.isdir(os.path.join(basedir,d))]
    for d in basedir.iterdir():
        if d.is_file():
            for filename in basedir.iterdir():
                path=os.path.join(basedir,filename)
                if path.endswith('.txt'):
                    with open(path,'r') as f:
                        data=f.read()
                        data=data.lower()
                        if text in data:
                            print(f"Found in text file: {path}")
                            break
                if path.endswith('.pdf'):
                    with open(path,'rb') as pdf:
                        pdfreader=pypdf2.PdfReader(pdf)
                        num=-1
                        for page in pdfreader.pages:
                            num=num+1
                            pdftext=pdfreader.pages[num].extract_text()
                            pdftext=pdftext.lower()
                            if text in str(pdftext):
                                print(f"Found in PDF file: {path}")
                                break
                if path.endswith('.docx'):
                    with open(path,'rb') as docx_file:
                        document_reader=docx.Document(docx_file)
                        fulltext=[]
                        for para in document_reader.paragraphs:
                            fulltext.append(para.text)
                            docx_text='\n'.join(fulltext)
                            docx_text=docx_text.lower()
                        if text in docx_text:
                            print(f"Found in DOCX file: {path}")
text="hydrogen"
basedir=Path(r"C:\avyan\python\IMPORTANT_MODULES")
t1=threading.Thread(target=search_text_in_directory(text,basedir))
t2=threading.Thread(target=search_text_in_directory_directory(text,basedir))
start = time.perf_counter()
t1.start()
t2.start()
t1.join()
t2.join()
end = time.perf_counter()
print(f"Elapsed time: {end - start:.4f} seconds")